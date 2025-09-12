import docx
import os
import shutil
import json
import re

# --- HELPER FUNCTIONS ---
def _iter_block_items(parent):
    if isinstance(parent, docx.document.Document): parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell): parent_elm = parent._tc
    else: raise ValueError("Parent must be a Document or _Cell object")
    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P): yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl): yield docx.table.Table(child, parent)

def load_word_doc_to_string(folder_path):
    filename=None;
    try:
        if os.path.isdir(folder_path):
            for f in os.listdir(folder_path):
                if f.lower().endswith('.docx') and not f.startswith('~$'):filename=os.path.join(folder_path, f);break
        elif os.path.isfile(folder_path):filename=folder_path
    except FileNotFoundError:return f"Error: Directory not found at '{folder_path}'"
    if not filename:return f"Error: No .docx file found in '{folder_path}'"
    try:
        doc=docx.Document(filename);blocks=[]
        for block in _iter_block_items(doc):
            if isinstance(block,docx.text.paragraph.Paragraph):
                if block.text.strip():blocks.append(block.text)
            elif isinstance(block,docx.table.Table):
                if not block.rows:continue
                lines=["| "+" | ".join(c.text.replace('\n',' ').strip() for c in block.rows[0].cells)+" |","| "+" | ".join(['---']*len(block.rows[0].cells))+" |"]
                for row in block.rows[1:]:lines.append("| "+" | ".join(c.text.replace('\n',' ').strip() for c in row.cells)+" |")
                blocks.append("\n".join(lines))
        return "\n\n".join(blocks)
    except Exception as e:return f"Error processing file '{os.path.basename(filename)}': {e}"

def create_output_doc_from_template(project_name):
    template_folder,output_folder="output_template","auto_pdd_output"
    template_path=next((os.path.join(template_folder,f) for f in os.listdir(template_folder) if f.lower().endswith('.docx') and not f.startswith('~$')),None)
    if not template_path:raise FileNotFoundError(f"Error: No .docx template found in '{template_folder}'")
    if not os.path.exists(output_folder):os.makedirs(output_folder)
    output_path=os.path.join(output_folder,f"AutoPDD_{project_name}.docx")
    if not os.path.exists(output_path):
        shutil.copy(template_path,output_path);print(f"Created output document at: {output_path}")
    else:
        print(f"Output document already exists at: {output_path}. This file will be updated.")
    return output_path

def get_ai_value(ai_json_data, text):
    """GENERIC: Get best matching AI value for any text"""
    # Handle both dict and string formats
    for key, data in ai_json_data.items():
        if isinstance(data, dict):
            value = data.get("value", "")
        elif isinstance(data, str):
            value = data
        else:
            continue
            
        if value and value != "INFO_NOT_FOUND":
            # Check if key matches text content
            key_words = set(key.lower().split())
            text_words = set(text.lower().split())
            
            # If there's word overlap, this might be a match
            if key_words.intersection(text_words):
                return value
    
    return None

def simple_replace_section(doc_path, start_marker, end_marker, ai_json_data, status):
    """ULTRA SIMPLE: Just fill empty things with AI content"""
    try:
        doc = docx.Document(doc_path)
        print(f"Processing: {start_marker}")
        
        # Find section boundaries
        in_section = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Check section boundaries
            if start_marker.split()[-1] in text:  # Match by section title
                in_section = True
                continue
            elif end_marker.split()[-1] in text:
                in_section = False
                break
                
            if not in_section:
                continue
                
            # GENERIC PROCESSING:
            
            # 1. CHECKBOXES: If has [ ] or [X], try to match with AI
            if '[' in text and ']' in text:
                ai_value = get_ai_value(ai_json_data, text)
                if ai_value:
                    # Check if AI value matches this option
                    text_lower = text.lower().replace('[', '').replace(']', '').replace('x', '').strip()
                    if any(word in ai_value.lower() for word in text_lower.split() if len(word) > 3):
                        paragraph.text = text.replace('[ ]', '[X]').replace('[x]', '[X]').replace('[X]', '[X]')
                        print(f"Selected: {text[:50]}...")
                    else:
                        paragraph.text = text.replace('[X]', '[ ]').replace('[x]', '[ ]')
                continue
            
            # 2. EMPTY/INSTRUCTION PARAGRAPHS: Add AI content 
            if (len(text) > 20 and 
                any(word in text.lower() for word in ['describe', 'provide', 'justify', 'explain']) and
                not text.startswith('•')):
                
                ai_value = get_ai_value(ai_json_data, text)
                if ai_value:
                    # Add AI content as new paragraph after this one
                    new_para = paragraph._element.getparent().insert(
                        list(paragraph._element.getparent()).index(paragraph._element) + 1,
                        doc.add_paragraph(ai_value)._element
                    )
                    print(f"Added content after: {text[:50]}...")
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # If cell is empty or placeholder, try to fill
                    if (not cell_text or 
                        len(cell_text) < 3 or
                        cell_text in ['...', 'TBD', 'N/A'] or
                        cell_text.startswith('(') and cell_text.endswith(')')):
                        
                        ai_value = get_ai_value(ai_json_data, cell_text)
                        if ai_value:
                            cell.text = ai_value
                            print(f"Filled cell: {ai_value[:30]}...")
                        else:
                            cell.text = "INFO_NOT_FOUND"
        
        doc.save(doc_path)
        print(f"Saved: {start_marker}")
        
    except Exception as e:
        print(f"ERROR: {e}")