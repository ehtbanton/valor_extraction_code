import docx
import os
import shutil
import json

# --- HELPER FUNCTIONS (UNCHANGED FROM ORIGINAL) ---
def _iter_block_items(parent):
    if isinstance(parent, docx.document.Document): parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell): parent_elm = parent._tc
    else: raise ValueError("Parent must be a Document or _Cell object")
    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P): yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl): yield docx.table.Table(child, parent)

def load_word_doc_to_string(folder_path):
    filename=None;
    try:
        if os.path.isdir(folder_path):
            for f in os.listdir(folder_path):
                if f.lower().endswith('.docx') and not f.startswith('~$'):filename=os.path.join(folder_path, f);break
        elif os.path.isfile(folder_path):filename=folder_path
    except FileNotFoundError:return f"Error: Directory not found at '{folder_path}'"
    if not filename:return f"Error: No .docx file found in '{folder_path}'"
    try:
        doc=docx.Document(filename);blocks=[]
        for block in _iter_block_items(doc):
            if isinstance(block,docx.text.paragraph.Paragraph):
                if block.text.strip():blocks.append(block.text)
            elif isinstance(block,docx.table.Table):
                if not block.rows:continue
                lines=["| "+" | ".join(c.text.replace('\n',' ').strip() for c in block.rows[0].cells)+" |","| "+" | ".join(['---']*len(block.rows[0].cells))+" |"]
                for row in block.rows[1:]:lines.append("| "+" | ".join(c.text.replace('\n',' ').strip() for c in row.cells)+" |")
                blocks.append("\n".join(lines))
        return "\n\n".join(blocks)
    except Exception as e:return f"Error processing file '{os.path.basename(filename)}': {e}"

def create_output_doc_from_template(project_name):
    template_folder,output_folder="output_template","auto_pdd_output"
    template_path=next((os.path.join(template_folder,f) for f in os.listdir(template_folder) if f.lower().endswith('.docx') and not f.startswith('~$')),None)
    if not template_path:raise FileNotFoundError(f"Error: No .docx template found in '{template_folder}'")
    if not os.path.exists(output_folder):os.makedirs(output_folder)
    output_path=os.path.join(output_folder,f"AutoPDD_{project_name}.docx")
    if not os.path.exists(output_path):
        shutil.copy(template_path,output_path);print(f"Created output document at: {output_path}")
    else:
        print(f"Output document already exists at: {output_path}. This file will be updated.")
    return output_path

def add_comment(paragraph, text, author="Source Reference"):
    try:
        # Only add a comment if there is actual sourced content.
        if paragraph.text.strip() and "INFO_NOT_FOUND" not in text and source.strip() != "N/A":
            paragraph.add_comment(text, author=author)
    except Exception:
        pass

def get_heading_level(paragraph):
    """Get the heading level (1-9) of a paragraph, or 0 if not a heading."""
    if paragraph.style and paragraph.style.name:
        style_name = paragraph.style.name.lower()
        if 'heading' in style_name:
            import re
            match = re.search(r'heading\s*(\d+)', style_name)
            if match:
                level = int(match.group(1))
                return level if 1 <= level <= 9 else 0
    return 0

# --- THE DEFINITIVE, WORKING SOLUTION ---

def replace_section_in_word_doc(doc_path, start_marker, end_marker, ai_json_data, status):
    try:
        doc = docx.Document(doc_path)
        all_blocks = list(_iter_block_items(doc))

        start_index, end_index = -1, len(all_blocks)
        start_heading_level = 0
        matches_found = []
        
        for i, block in enumerate(all_blocks):
            if isinstance(block, docx.text.paragraph.Paragraph):
                text = block.text.strip()
                current_heading_level = get_heading_level(block)
                
                # Find start marker - avoid TOC entries (they have numbers and tabs)
                if (text == start_marker or 
                    (len(start_marker.split()) > 1 and start_marker in text and
                     not ('\t' in text and any(char.isdigit() for char in text.split('\t')[-1])))):
                    matches_found.append(f"Block {i}: '{text}' (Level {current_heading_level})")
                    # Only use the FIRST valid match to avoid loop issues
                    if start_index == -1:
                        start_index = i
                        start_heading_level = current_heading_level
                        print(f"  > Selected start section '{text}' at heading level {start_heading_level} (Block {i})")
                    
                # Find end marker - heading level aware
                elif start_index != -1:
                    # Exact or fuzzy end marker match
                    if (text == end_marker or 
                        (len(end_marker.split()) > 1 and end_marker in text)):
                        end_index = i
                        print(f"  > Found end marker '{text}'")
                        break
                    # Or if we hit a heading at same or higher level (and we started at a heading)
                    elif (start_heading_level > 0 and current_heading_level > 0 and 
                          current_heading_level <= start_heading_level):
                        end_index = i
                        print(f"  > Found end boundary: higher level heading (level {current_heading_level})")
                        break
        
        if start_index == -1:
            print(f"  > WARNING: Start marker '{start_marker}' not found.")
            if matches_found:
                print(f"  > Potential matches found: {matches_found}")
            return
        else:
            if len(matches_found) > 1:
                print(f"  > Multiple matches found: {matches_found}")
                print(f"  > Using first valid match: Block {start_index}")

        # Fix for section skipping: Find or insert the status paragraph correctly.
        status_p_index = start_index + 1
        if status_p_index < len(all_blocks) and isinstance(all_blocks[status_p_index], docx.text.paragraph.Paragraph) and "SECTION_" in all_blocks[status_p_index].text:
            all_blocks[status_p_index].text = status
        else:
            all_blocks[start_index].insert_paragraph_before(status)
        print(f"  > Status for '{start_marker}' set to: {status}")

        section_blocks = all_blocks[start_index : end_index]

        for block in section_blocks:
            # Handle paragraphs with enhanced key matching for hierarchical content
            if isinstance(block, docx.text.paragraph.Paragraph):
                block_text = block.text.strip()
                if not block_text:
                    continue
                
                # Skip headings - don't replace heading text with AI content
                if get_heading_level(block) > 0:
                    continue
                    
                best_match_key = None
                best_match_score = 0
                
                for key, data_obj in ai_json_data.items():
                    if not isinstance(data_obj, dict):
                        continue
                        
                    # Enhanced matching for hierarchical AI keys
                    score = 0
                    
                    # Method 1: Direct substring match
                    if key in block_text:
                        score = 100
                    # Method 2: Extract content after the last dash/colon and match
                    elif ' - ' in key:
                        key_content = key.split(' - ')[-1]  # Get part after last dash
                        if key_content in block_text:
                            score = 90
                    elif ': ' in key:
                        # Try both prefix and suffix matching
                        key_parts = key.split(': ')
                        key_prefix = key_parts[0].strip()  # Part before colon
                        
                        # Check if prefix matches (e.g., "General eligibility" in document)
                        if key_prefix in block_text:
                            score = 85
                        # Check if suffix matches
                        elif len(key_parts) > 1 and key_parts[-1] in block_text:
                            score = 80
                    # Method 3: Word overlap scoring
                    else:
                        key_words = set(w.lower() for w in key.split() if len(w) > 3)
                        block_words = set(w.lower() for w in block_text.split() if len(w) > 3)
                        if key_words and block_words:
                            overlap = len(key_words.intersection(block_words))
                            if overlap > 0:
                                score = (overlap / len(key_words)) * 80
                    
                    if score > best_match_score:
                        best_match_score = score
                        best_match_key = key
                
                # Apply the best match if score is high enough, OR fill with INFO_NOT_FOUND if no match
                if best_match_key and best_match_score >= 70:
                    value = ai_json_data[best_match_key].get("value", "")
                    source = ai_json_data[best_match_key].get("source", "N/A")
                    
                    # Process ALL content - whether it has value or INFO_NOT_FOUND
                    if value:
                        if value != "INFO_NOT_FOUND":
                            # Special handling for checkbox content - detect by indentation and context
                            is_checkbox_option = (
                                block_text.startswith('☐') or block_text.startswith('[ ]') or block_text.startswith('[X]') or
                                (block_text.startswith('  ') and any(option in block_text for option in ['location', 'grouped', 'project']))
                            )
                            
                            if is_checkbox_option:
                                # This is a checkbox - handle Yes/No responses
                                if value.lower() in ['yes', 'y', 'true', '1']:
                                    # Check the box and keep the original text
                                    original_text = block_text.replace('☐', '').replace('[ ]', '').replace('[X]', '').strip()
                                    block.text = f"☒ {original_text}"
                                    add_comment(block, f"Checked based on: {source}")
                                    print(f"  > Checked checkbox (score {best_match_score:.0f}): {best_match_key[:50]}...")
                                elif value.lower() in ['no', 'n', 'false', '0']:
                                    # Leave unchecked but ensure proper formatting
                                    original_text = block_text.replace('☐', '').replace('[ ]', '').replace('[X]', '').strip()
                                    block.text = f"☐ {original_text}"
                                    add_comment(block, f"Unchecked based on: {source}")
                                    print(f"  > Left checkbox unchecked (score {best_match_score:.0f}): {best_match_key[:50]}...")
                                else:
                                    # Non-boolean response for checkbox - replace text
                                    block.text = value
                                    add_comment(block, f"Source: {source}")
                                    print(f"  > Filled checkbox content (score {best_match_score:.0f}): {best_match_key[:50]}...")
                            else:
                                # Regular text replacement
                                block.text = value
                                add_comment(block, f"Source: {source}")
                                print(f"  > Filled hierarchical content (score {best_match_score:.0f}): {best_match_key[:50]}...")
                        else:
                            # Handle INFO_NOT_FOUND case
                            key_description = best_match_key.split(': ')[-1] if ': ' in best_match_key else best_match_key
                            block.text = f"INFO_NOT_FOUND: {key_description}"
                            add_comment(block, f"Information not found in source documents")
                            print(f"  > Filled with INFO_NOT_FOUND (score {best_match_score:.0f}): {best_match_key[:50]}...")
                
                # Handle blocks that don't match any AI key (score < 70)
                elif block_text and not block_text.startswith("INFO_NOT_FOUND"):
                    # This is content that doesn't match any AI key - mark as not found
                    block.text = f"INFO_NOT_FOUND: {block_text[:100]}"
                    add_comment(block, f"No matching AI response found for this content")
                    print(f"  > Filled with INFO_NOT_FOUND (no match): {block_text[:50]}...")

            # Handle tables with intelligent key matching
            elif isinstance(block, docx.table.Table):
                headers = [h.text.strip() for h in block.rows[0].cells]
                
                # Logic for simple 2-column key-value tables (e.g., Sectoral Scope)
                if len(headers) == 2:
                    for row in block.rows:
                        label_cell = row.cells[0]
                        value_cell = row.cells[1]
                        label_text = label_cell.text.strip()
                        if label_text in ai_json_data:
                            data = ai_json_data[label_text]
                            value = data.get("value", "")
                            source = data.get("source", "N/A")
                            # Fill with actual value or INFO_NOT_FOUND message
                            if value == "INFO_NOT_FOUND":
                                value_cell.text = f"INFO_NOT_FOUND: {label_text}"
                            else:
                                value_cell.text = value
                            add_comment(value_cell.paragraphs[0], f"Source: {source}")
                        else:
                            # No AI data for this table row - mark as not found
                            value_cell.text = f"INFO_NOT_FOUND: {label_text}"
                            add_comment(value_cell.paragraphs[0], f"No AI response found for this field")
                
                # Logic for multi-column data tables (e.g., Audit History)
                else:
                    for row in block.rows[1:]: # Skip header row
                        # Get the primary context from the first column of the row
                        row_context = row.cells[0].text.strip()
                        if not row_context or "..." in row_context: continue

                        # Clean up context for better matching (e.g., "Validation/verification" -> "Validation")
                        clean_row_context = row_context.split('/')[0]

                        for i, cell in enumerate(row.cells):
                            # Skip the first cell as it's our context
                            if i == 0: continue
                            
                            header = headers[i]
                            
                            # Find the best matching key from the AI's flat dictionary
                            best_match_key = None
                            for key in ai_json_data:
                                # A key is a match if it contains both the row context and the column header
                                if clean_row_context in key and header in key:
                                    best_match_key = key
                                    break
                            
                            if best_match_key:
                                data = ai_json_data[best_match_key]
                                value = data.get("value", "")
                                source = data.get("source", "N/A")
                                # Fill with actual value or INFO_NOT_FOUND message
                                if value == "INFO_NOT_FOUND":
                                    cell.text = f"INFO_NOT_FOUND: {clean_row_context} - {header}"
                                else:
                                    cell.text = value
                                if cell.paragraphs:
                                    add_comment(cell.paragraphs[0], f"Source: {source}")
                            else:
                                # No matching AI key for this table cell - mark as not found
                                cell.text = f"INFO_NOT_FOUND: {clean_row_context} - {header}"
                                if cell.paragraphs:
                                    add_comment(cell.paragraphs[0], f"No AI response found for this field")

        doc.save(doc_path)
        print(f"Successfully updated section '{start_marker}' in {os.path.basename(doc_path)}.")

    except Exception as e:
        print(f"FATAL ERROR during document generation for section '{start_marker}': {e}")