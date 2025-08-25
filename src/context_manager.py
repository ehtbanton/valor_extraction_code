import os
import json
import pdfplumber
import docx

def _extract_text_from_file(file_path):
    """
    A helper function to extract text and tables from a single file.
    
    Args:
        file_path (str): The full path to the .pdf or .docx file.
    
    Returns:
        str: The extracted text content, with tables in Markdown format.
             Returns an empty string if the file cannot be processed.
    """
    content_parts = []
    filename = os.path.basename(file_path)

    try:
        # --- Handle PDF files ---
        if filename.lower().endswith('.pdf'):
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        content_parts.append(page_text)
                    
                    # Extract tables and convert to Markdown
                    tables = page.extract_tables()
                    for table in tables:
                        if not table: continue
                        header = "| " + " | ".join(str(cell) if cell is not None else '' for cell in table[0]) + " |"
                        separator = "| " + " | ".join(["---"] * len(table[0])) + " |"
                        rows = ["| " + " | ".join(str(cell) if cell is not None else '' for cell in row) + " |" for row in table[1:]]
                        markdown_table = "\n".join([header, separator] + rows)
                        content_parts.append(f"\n\n--- Table on Page {i+1} ---\n{markdown_table}\n")

        # --- Handle Word (.docx) files ---
        elif filename.lower().endswith('.docx'):
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                content_parts.append(para.text)
            
            # Extract tables and convert to Markdown
            for i, table in enumerate(doc.tables):
                if not table.rows: continue
                header_cells = table.rows[0].cells
                header = "| " + " | ".join(cell.text.strip() for cell in header_cells) + " |"
                separator = "| " + " | ".join(["---"] * len(header_cells)) + " |"
                rows = ["| " + " | ".join(cell.text.strip() for cell in row.cells) + " |" for row in table.rows[1:]]
                markdown_table = "\n".join([header, separator] + rows)
                content_parts.append(f"\n\n--- Table {i+1} ---\n{markdown_table}\n")

    except Exception as e:
        print(f"Could not process file '{filename}'. Reason: {e}")
        return "" # Return empty string on failure

    return "\n".join(content_parts)


def extract_text_from_folder(folder_path):
    """
    Extracts text from PDF and Word files in a folder and maintains a TXT
    file containing the content in a structured (JSON) format, updating it 
    with any new or deleted files.

    Args:
        folder_path (str): The absolute or relative path to the folder.

    Returns:
        bool: True if the TXT file was modified (files added/removed),
              False otherwise.
    """
    if not os.path.isdir(folder_path):
        print(f"Error: Folder not found at '{folder_path}'")
        return False

    # The only change needed is the file extension
    txt_filepath = os.path.join(folder_path, "all_context.txt")
    changes_made = False

    # 1. Load existing data from all_context.txt or create an empty list
    try:
        with open(txt_filepath, 'r', encoding='utf-8') as f:
            # Read the file's text content, then parse it as JSON
            all_context = json.loads(f.read())
    except (FileNotFoundError, json.JSONDecodeError):
        all_context = []
    
    # Get a set of filenames we already have processed
    known_files = {entry['filename'] for entry in all_context}

    # Get a set of current .pdf and .docx files in the folder
    current_files = {
        f for f in os.listdir(folder_path) 
        if f.lower().endswith(('.pdf', '.docx'))
    }

    # 2. Handle deletions
    files_to_remove = known_files - current_files
    if files_to_remove:
        print(f"Files removed: {', '.join(files_to_remove)}")
        all_context = [
            entry for entry in all_context 
            if entry['filename'] not in files_to_remove
        ]
        changes_made = True

    # 3. Handle additions
    files_to_add = current_files - known_files
    if files_to_add:
        print(f"New files found: {', '.join(files_to_add)}")
        for filename in files_to_add:
            file_path = os.path.join(folder_path, filename)
            print(f"-> Processing: {filename}")
            
            text_content = _extract_text_from_file(file_path)
            
            if text_content:
                all_context.append({
                    'filename': filename,
                    'text_content': text_content
                })
                print(f"   ...extracted {len(text_content)} characters.")
                changes_made = True

    # 4. Save the updated data back to the TXT file if any changes were made
    if changes_made:
        print(f"\nSaving changes to '{txt_filepath}'...")
        try:
            with open(txt_filepath, 'w', encoding='utf-8') as f:
                # Convert the Python list to a JSON-formatted string
                json_string = json.dumps(all_context, indent=4)
                # Write that string to the .txt file
                f.write(json_string)
            print("...Success!")
        except Exception as e:
            print(f"Error saving the TXT file: {e}")
            return False
    else:
        print("\nNo changes detected. Content is up-to-date.")

    return changes_made