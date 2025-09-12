import docx
import os
import re
from word_editor import get_value_for_key_pattern

def replace_content_robustly(doc_path, start_marker, end_marker, ai_json_data, status):
    """
    A completely new approach to document processing that:
    1. Maps document structure explicitly
    2. Uses position-based content assignment
    3. Handles subsections properly
    4. Ensures all table cells are filled
    """
    try:
        doc = docx.Document(doc_path)
        print(f"=== ROBUST PROCESSING: {start_marker} ===")
        
        # Step 1: Map the entire document structure
        doc_map = analyze_document_structure(doc, start_marker, end_marker)
        
        # Step 2: Process content based on explicit mapping
        process_mapped_content(doc, doc_map, ai_json_data, start_marker)
        
        doc.save(doc_path)
        print(f"✓ Robust processing complete for {start_marker}")
        
    except Exception as e:
        print(f"✗ Robust processing failed: {e}")
        import traceback
        traceback.print_exc()

def analyze_document_structure(doc, start_marker, end_marker):
    """Analyze and map the document structure explicitly"""
    doc_map = {
        'section_start': -1,
        'section_end': -1,
        'subsections': {},
        'tables': [],
        'paragraphs': [],
        'bullets': []
    }
    
    in_section = False
    current_subsection = None
    
    # Extract just the title for matching
    start_title = start_marker.split(' ', 1)[1] if ' ' in start_marker else start_marker
    end_title = end_marker.split(' ', 1)[1] if ' ' in end_marker else end_marker
    
    # Map paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Find section boundaries
        if text == start_title:
            in_section = True
            doc_map['section_start'] = i
            print(f"  📍 Section starts at paragraph {i}: {text}")
            continue
        elif text == end_title:
            in_section = False
            doc_map['section_end'] = i
            print(f"  📍 Section ends at paragraph {i}: {text}")
            break
            
        if not in_section:
            continue
            
        # Detect subsections with regex
        subsection_match = re.match(r'^(\d+\.\d+\.\d+)\s+(.+)$', text)
        if subsection_match:
            current_subsection = subsection_match.group(1)
            subsection_title = subsection_match.group(2)
            doc_map['subsections'][current_subsection] = {
                'paragraph_index': i,
                'title': subsection_title,
                'content_paragraphs': []
            }
            print(f"  🎯 Found subsection {current_subsection}: {subsection_title}")
            continue
            
        # Map paragraph content to current subsection
        if current_subsection and text:
            doc_map['subsections'][current_subsection]['content_paragraphs'].append(i)
            
        # Map bullets
        if text.startswith('•') or text == '•' or text == '':
            doc_map['bullets'].append(i)
            
        # Map regular paragraphs
        if text and not text.startswith('•'):
            doc_map['paragraphs'].append({
                'index': i,
                'text': text[:100],
                'subsection': current_subsection
            })
    
    # Map tables in section
    table_count = 0
    for i, block in enumerate(_iter_block_items(doc)):
        if isinstance(block, docx.table.Table):
            # Check if table is in our section (rough approximation)
            if doc_map['section_start'] <= (i * 2) <= (doc_map['section_end'] * 2 if doc_map['section_end'] > 0 else len(doc.paragraphs) * 2):
                headers = [cell.text.strip() for cell in block.rows[0].cells] if block.rows else []
                doc_map['tables'].append({
                    'table_obj': block,
                    'headers': headers,
                    'index': table_count
                })
                table_count += 1
                print(f"  📊 Found table {table_count}: {headers}")
    
    return doc_map

def process_mapped_content(doc, doc_map, ai_json_data, start_marker):
    """Process content using explicit document mapping"""
    
    # Process subsections with different content
    subsection_processors = {
        '1.4.1': process_general_eligibility,
        '1.4.2': process_afolu_eligibility, 
        '1.4.3': process_transfer_eligibility
    }
    
    for subsection_id, subsection_data in doc_map['subsections'].items():
        if subsection_id in subsection_processors:
            processor = subsection_processors[subsection_id]
            processor(doc, subsection_data, ai_json_data)
    
    # Process tables with explicit mapping
    process_tables_explicitly(doc_map['tables'], ai_json_data, start_marker)
    
    # Process bullets with position-based assignment
    process_bullets_explicitly(doc, doc_map['bullets'], ai_json_data, start_marker)

def process_general_eligibility(doc, subsection_data, ai_json_data):
    """Process 1.4.1 General eligibility with specific content"""
    print("  🎯 Processing 1.4.1 General eligibility")
    
    value = get_value_for_key_pattern(ai_json_data, ["paragraph: general eligibility", "general eligibility"])
    if value and value != "INFO_NOT_FOUND":
        # Fill only the first paragraph in this subsection
        if subsection_data['content_paragraphs']:
            para_index = subsection_data['content_paragraphs'][0]
            doc.paragraphs[para_index].text = value
            print(f"    ✓ Filled paragraph {para_index} with general eligibility content")
            
            # Clear other paragraphs in this subsection to prevent duplication
            for i in subsection_data['content_paragraphs'][1:]:
                doc.paragraphs[i].text = ""
                print(f"    ✓ Cleared paragraph {i}")

def process_afolu_eligibility(doc, subsection_data, ai_json_data):
    """Process 1.4.2 AFOLU eligibility with specific content"""
    print("  🎯 Processing 1.4.2 AFOLU eligibility")
    
    value = get_value_for_key_pattern(ai_json_data, ["paragraph: afolu project eligibility", "afolu"])
    if not value or value == "INFO_NOT_FOUND":
        value = "This project is not an AFOLU (Agriculture, Forestry and Other Land Use) project. The information requested in this section is not applicable."
    
    if subsection_data['content_paragraphs']:
        para_index = subsection_data['content_paragraphs'][0]
        doc.paragraphs[para_index].text = value
        print(f"    ✓ Filled paragraph {para_index} with AFOLU content")
        
        # Clear other paragraphs
        for i in subsection_data['content_paragraphs'][1:]:
            doc.paragraphs[i].text = ""

def process_transfer_eligibility(doc, subsection_data, ai_json_data):
    """Process 1.4.3 Transfer eligibility with specific content"""
    print("  🎯 Processing 1.4.3 Transfer eligibility")
    
    value = get_value_for_key_pattern(ai_json_data, ["paragraph: transfer project eligibility", "transfer"])
    if not value or value == "INFO_NOT_FOUND":
        value = "This project is not a transfer project. The information requested in this section is not applicable."
    
    if subsection_data['content_paragraphs']:
        para_index = subsection_data['content_paragraphs'][0]
        doc.paragraphs[para_index].text = value
        print(f"    ✓ Filled paragraph {para_index} with transfer content")
        
        # Clear other paragraphs
        for i in subsection_data['content_paragraphs'][1:]:
            doc.paragraphs[i].text = ""

def process_tables_explicitly(tables, ai_json_data, start_marker):
    """Process tables with explicit cell-by-cell mapping"""
    print(f"  📊 Processing {len(tables)} tables explicitly")
    
    for table_info in tables:
        table = table_info['table_obj']
        headers = table_info['headers']
        
        print(f"    📋 Table with headers: {headers}")
        
        for row_idx in range(1, len(table.rows)):  # Skip header row
            row = table.rows[row_idx]
            row_label = row.cells[0].text.strip() if row.cells else ""
            
            for col_idx in range(1, len(row.cells)):  # Skip label column
                cell = row.cells[col_idx]
                col_header = headers[col_idx] if col_idx < len(headers) else ""
                
                # Check if cell needs filling
                if not cell.text.strip() or cell.text.strip() in ['', ' ', '...', '…']:
                    filled = False
                    
                    # Explicit sectoral scope handling
                    if row_label.lower() == "sectoral scope":
                        value = get_value_for_key_pattern(ai_json_data, [
                            "table: sectoral scope - non-afolu", 
                            "table: sectoral scope", 
                            "sectoral scope", 
                            "energy"
                        ])
                        if value and value != "INFO_NOT_FOUND":
                            cell.text = value
                            print(f"      ✓ Filled sectoral scope: {value}")
                            filled = True
                    
                    # Project activity type
                    elif "project activity type" in row_label.lower():
                        value = get_value_for_key_pattern(ai_json_data, [
                            "table: project activity type - non-afolu",
                            "project activity type", 
                            "renewable energy"
                        ])
                        if value and value != "INFO_NOT_FOUND":
                            cell.text = value
                            print(f"      ✓ Filled project activity: {value}")
                            filled = True
                    
                    # Generic fallback with better matching
                    if not filled:
                        search_terms = []
                        if row_label:
                            search_terms.extend([row_label, f"table: {row_label.lower()}"])
                        if col_header:
                            search_terms.extend([col_header, f"table: {col_header.lower()}"])
                        
                        value = get_value_for_key_pattern(ai_json_data, search_terms)
                        if value and value != "INFO_NOT_FOUND":
                            cell.text = value
                            print(f"      ✓ Filled {row_label}/{col_header}: {value}")
                            filled = True
                    
                    # Ensure no cell is left empty
                    if not filled:
                        missing_info = f"{row_label}" + (f", {col_header}" if col_header else "")
                        cell.text = f"INFO_NOT_FOUND: {missing_info.lower()}"
                        print(f"      ✓ Default: {missing_info}")

def process_bullets_explicitly(doc, bullet_indices, ai_json_data, start_marker):
    """Process bullets with explicit position-based assignment"""
    print(f"  🔸 Processing {len(bullet_indices)} bullets explicitly")
    
    bullet_content_map = [
        ("summary", ["paragraph: a summary description", "summary", "technologies"]),
        ("location", ["paragraph: the location", "location", "kampong"]),
        ("ghg", ["paragraph: an explanation", "ghg", "emission", "reductions"]),
        ("prior", ["paragraph: a brief description", "prior", "scenario"]),
        ("estimate", ["paragraph: an estimate", "estimate", "annual", "reductions"])
    ]
    
    for i, para_idx in enumerate(bullet_indices):
        if i < len(bullet_content_map):
            bullet_name, search_patterns = bullet_content_map[i]
            value = get_value_for_key_pattern(ai_json_data, search_patterns)
            
            if value and value != "INFO_NOT_FOUND":
                # Use first sentence for bullet
                first_sentence = value.split('.')[0].strip()
                if not first_sentence.endswith('.'):
                    first_sentence += '.'
                doc.paragraphs[para_idx].text = f"• {first_sentence}"
                print(f"    ✓ Bullet {i+1} ({bullet_name}): {first_sentence[:50]}...")
            else:
                # Fallback content
                fallbacks = [
                    "Summary description of technologies/measures implemented by the project.",
                    "Location of the project.",
                    "How the project generates GHG emission reductions.",
                    "Scenario existing prior to project implementation.",
                    "Estimate of annual average and total reductions."
                ]
                doc.paragraphs[para_idx].text = f"• {fallbacks[i]}"
                print(f"    ✓ Bullet {i+1} fallback: {fallbacks[i]}")

def _iter_block_items(parent):
    """Helper function from original code"""
    if isinstance(parent, docx.document.Document): 
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell): 
        parent_elm = parent._tc
    else: 
        raise ValueError("Parent must be a Document or _Cell object")
    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P): 
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl): 
            yield docx.table.Table(child, parent)