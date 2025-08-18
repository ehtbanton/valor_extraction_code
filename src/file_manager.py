
import os
import pdfplumber
import docx

def extract_text_from_folder(folder_path):
    """
    Extracts text and tables from all PDF and Word (.docx) files in a folder,
    combines all content into a single string, saves it to a single .txt file
    at 'folder_path/combined_text/provided_docs.txt', and returns the
    combined string. Tables are converted into Markdown format.

    This function will first check if the output file already exists to avoid
    unnecessary processing.

    Args:
        folder_path (str): The absolute or relative path to the folder
                           containing the files.

    Returns:
        str: A single string containing all the extracted text content
             (including Markdown tables) from all files, separated by
             two newlines. Returns an empty string if the folder does not
             exist, is empty, or if the output file already exists.
    """
    # Check if the provided path is a valid directory
    if not os.path.isdir(folder_path):
        print(f"Error: Folder not found at '{folder_path}'")
        return ""

    # Define the output directory and filename
    output_filepath = os.path.join(folder_path, "provided_docs.txt")

    # Check if the output file already exists to avoid re-processing
    if os.path.exists(output_filepath):
        print(f"Output file already exists at '{output_filepath}'. Skipping extraction.")
        # Optionally, you could read and return the content of the existing file here
        # with open(output_filepath, 'r', encoding='utf-8') as f:
        #     return f.read()
        return ""

    all_content_parts = []
    print(f"Scanning folder: {folder_path}")

    # Iterate over each file in the specified directory
    for filename in os.listdir(folder_path):
        # Construct the full file path
        file_path = os.path.join(folder_path, filename)

        # Skip subdirectories and the output directory itself
        if not os.path.isfile(file_path):
            continue

        current_file_content = []
        try:
            # --- Handle PDF files ---
            if filename.lower().endswith('.pdf'):
                print(f"-> Processing PDF: {filename}")
                with pdfplumber.open(file_path) as pdf:
                    # Loop through each page
                    for i, page in enumerate(pdf.pages):
                        # Extract plain text from the page
                        page_text = page.extract_text()
                        if page_text:
                            current_file_content.append(page_text)

                        # --- Extract tables and convert to Markdown ---
                        tables = page.extract_tables()
                        for table in tables:
                            if not table: continue

                            # Build the Markdown table header
                            header = "| " + " | ".join(str(cell) if cell is not None else '' for cell in table[0]) + " |"
                            separator = "| " + " | ".join(["---"] * len(table[0])) + " |"

                            # Build the Markdown table rows
                            rows = []
                            for row in table[1:]:
                                rows.append("| " + " | ".join(str(cell) if cell is not None else '' for cell in row) + " |")

                            markdown_table = "\n".join([header, separator] + rows)
                            current_file_content.append(f"\n\n--- Table on Page {i+1} ---\n{markdown_table}\n")

            # --- Handle Word (.docx) files ---
            elif filename.lower().endswith('.docx'):
                print(f"-> Processing DOCX: {filename}")
                doc = docx.Document(file_path)

                # Loop through each paragraph and extract text
                for para in doc.paragraphs:
                    current_file_content.append(para.text)

                # --- Extract tables and convert to Markdown ---
                for i, table in enumerate(doc.tables):
                    if not table.rows: continue

                    # Build the Markdown table header from the first row
                    header_cells = table.rows[0].cells
                    header = "| " + " | ".join(cell.text.strip() for cell in header_cells) + " |"
                    separator = "| " + " | ".join(["---"] * len(header_cells)) + " |"

                    # Build the Markdown table rows from the rest of the table
                    rows = []
                    for row in table.rows[1:]:
                        rows.append("| " + " | ".join(cell.text.strip() for cell in row.cells) + " |")

                    markdown_table = "\n".join([header, separator] + rows)
                    current_file_content.append(f"\n\n--- Table {i+1} ---\n{markdown_table}\n")

            # If content was extracted from the file, join it and add to the main list
            if current_file_content:
                final_text_for_file = "\n".join(current_file_content)
                all_content_parts.append(final_text_for_file)
                print(f"   ...extracted {len(final_text_for_file)} characters.")

        except Exception as e:
            # Handle potential errors with file processing
            print(f"Could not process file '{filename}'. Reason: {e}")

    # --- Combine all extracted content and save to a single file ---
    if not all_content_parts:
        print("No content was extracted from any files.")
        return ""

    # Join the content from all files with two newlines
    combined_text = "\n\n".join(all_content_parts)

    try:
        # Save the combined content to the single output file
        with open(output_filepath, 'w', encoding='utf-8') as output_file:
            output_file.write(combined_text)
        print(f"\nSuccess! All content combined and saved to '{output_filepath}'")
    except Exception as e:
        print(f"\nError saving the combined file: {e}")

    return combined_text


def name_files_in_folder(folder_path):
    provided_files_list = os.listdir(folder_path)
    for file in provided_files_list:
        provided_files_list[provided_files_list.index(file)] = folder_path + "/" + file
    return provided_files_list













